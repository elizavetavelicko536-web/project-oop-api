import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

from .parsers import read_file
from .core import AnnotationResult


class Annotator:
    def __init__(self, provider):
        self.provider = provider

    def __call__(self, files, prompts):
        results = []
        for i, f in enumerate(files):
            prompt = prompts[i] if i < len(prompts) else prompts[-1] if prompts else ""
            text = read_file(f)
            annotated = self.provider.annotate(text, prompt)

            result = AnnotationResult(
                file_name=f,
                prompt=prompt,
                text=annotated
            )
            results.append(result)

        return results  # ← Теперь возвращаем список AnnotationResult


def save_results_to_docx(results, out_dir="out"):
    out_dir = str(out_dir)

    if out_dir.endswith(".docx"):
        folder = os.path.dirname(out_dir) or "."
        filename = os.path.basename(out_dir)
    else:
        folder = out_dir or "out"
        filename = "annotations.docx"

    if not os.path.exists(folder):
        os.makedirs(folder)

    out_path = os.path.join(folder, filename)
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    if isinstance(results, dict):
        items = list(results.items())
    else:
        items = [(r.file_name, r.text) for r in results]

    for i, (file_name, annotation) in enumerate(items, 1):
        display_name = os.path.basename(file_name)

        heading = doc.add_heading(level=1)
        run = heading.add_run(f"{i}. {display_name}")
        run.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()

        for paragraph in annotation.split("\n"):
            if paragraph.strip():
                p = doc.add_paragraph(paragraph.strip())
                p.paragraph_format.space_after = Pt(6)

        if i < len(items):
            doc.add_paragraph("\n" + "=" * 50 + "\n")

    doc.save(out_path)
    return out_path