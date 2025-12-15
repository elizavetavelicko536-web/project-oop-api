from src.annotations.annotator import Annotator
from src.annotations.providers import MockProvider
from src.annotations.core import AnnotationResult
from docx import Document


def create_test_docx(path, text):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)


def test_annotator_returns_annotation_results(tmp_path):
    file_path = tmp_path / "test.docx"
    create_test_docx(file_path, "Это тестовый текст.")
    provider = MockProvider()
    annotator = Annotator(provider)
    results = annotator([str(file_path)], ["Сделай аннотацию"])

    assert isinstance(results, list)
    assert len(results) == 1
    assert isinstance(results[0], AnnotationResult)
    assert results[0].file_name == str(file_path)
    assert "MOCK RESPONSE" in results[0].text


def test_annotation_result_operations():
    r1 = AnnotationResult("file1.docx", "промт1", "текст1")
    r2 = AnnotationResult("file2.docx", "промт2", "текст2")

    r3 = r1 + r2
    assert r3.file_name == "file1.docx|file2.docx"
    assert r3.prompt == "промт1|промт2"
    assert r3.text == "текст1\n\nтекст2"

    r1 += r2
    assert r1.file_name == "file1.docx|file2.docx"
    assert r1.text == "текст1\n\nтекст2"
    assert r1.word_count == 2
    assert "file1.docx" in r1.file_basename


def test_save_results_with_annotation_results(tmp_path):
    from src.annotations.annotator import save_results_to_docx

    results = [
        AnnotationResult("file1.docx", "промт1", "Аннотация 1"),
        AnnotationResult("file2.docx", "промт2", "Аннотация 2"),
    ]
    out_dir = tmp_path / "out"
    out_path = save_results_to_docx(results, str(out_dir))

    assert out_path.endswith("annotations.docx")

    doc = Document(out_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "file1.docx" in text
    assert "file2.docx" in text
    assert "Аннотация 1" in text
    assert "Аннотация 2" in text